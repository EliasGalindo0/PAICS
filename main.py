import io
import os
import sys

import fitz  # PyMuPDF
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from dotenv import load_dotenv
from PIL import Image

from ai.analyzer import VetAIAnalyzer, load_dicom_image

load_dotenv()
if os.name == "nt":
    for stream_name in ("stdout", "stderr"):
        stream = getattr(sys, stream_name, None)
        if stream and hasattr(stream, "reconfigure"):
            try:
                stream.reconfigure(encoding="utf-8", errors="replace")
            except Exception:
                pass


def _safe_print(*args, **kwargs):
    """Print que não quebra em consoles Windows com encoding legado."""
    try:
        print(*args, **kwargs)
    except UnicodeEncodeError:
        sep = kwargs.get("sep", " ")
        end = kwargs.get("end", "\n")
        text = sep.join(str(a) for a in args) + end
        enc = getattr(sys.stdout, "encoding", None) or "utf-8"
        try:
            sys.stdout.write(text.encode(enc, errors="replace").decode(enc, errors="replace"))
        except Exception:
            sys.stdout.write(
                text.encode("utf-8", errors="replace").decode("utf-8", errors="replace")
            )


OUTPUT_DIR = os.getenv("OUTPUT_DIR", "laudos_com_ia")
PDF_ZOOM_FACTOR = float(os.getenv("PDF_ZOOM_FACTOR", "2.0"))
IMAGE_WIDTH_INCHES = float(os.getenv("IMAGE_WIDTH_INCHES", "5.5"))


class VetReportGenerator:
    def __init__(self, output_dir: str = None):
        self.output_dir = output_dir or OUTPUT_DIR
        self.ai_analyzer = VetAIAnalyzer()
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)

    def _load_images_from_path(self, path: str) -> list:
        """Carrega imagens de um arquivo ou diretório"""
        images = []

        if os.path.isfile(path):
            # Arquivo único
            ext = os.path.splitext(path)[1].lower()

            if ext == '.pdf':
                images = self._pdf_to_pil_images(path)
            elif ext in ['.dcm', '.dicom']:
                img = load_dicom_image(path)
                if img:
                    images.append(img)
            elif ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff']:
                try:
                    img = Image.open(path)
                    if img.mode != 'RGB':
                        img = img.convert('RGB')
                    images.append(img)
                except Exception as e:
                    print(f"Erro ao carregar {path}: {e}")

        elif os.path.isdir(path):
            # Diretório com múltiplas imagens
            valid_extensions = ['.pdf', '.jpg', '.jpeg', '.png', '.dcm', '.dicom', '.bmp', '.tiff']

            for filename in sorted(os.listdir(path)):
                file_path = os.path.join(path, filename)
                ext = os.path.splitext(filename)[1].lower()

                if ext in valid_extensions:
                    if ext == '.pdf':
                        images.extend(self._pdf_to_pil_images(file_path))
                    elif ext in ['.dcm', '.dicom']:
                        img = load_dicom_image(file_path)
                        if img:
                            images.append(img)
                    elif ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff']:
                        try:
                            img = Image.open(file_path)
                            if img.mode != 'RGB':
                                img = img.convert('RGB')
                            images.append(img)
                        except Exception as e:
                            print(f"Erro ao carregar {file_path}: {e}")

        return images

    def _pdf_to_pil_images(self, pdf_path: str) -> list:
        """Converte páginas do PDF em objetos PIL.Image"""
        doc = fitz.open(pdf_path)
        pil_images = []

        for page in doc:
            zoom = PDF_ZOOM_FACTOR
            pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom))
            mode = "RGBA" if pix.alpha else "RGB"
            img = Image.frombytes(mode, [pix.width, pix.height], pix.samples)
            pil_images.append(img)

        doc.close()
        return pil_images

    def create_report(self, path: str):
        """Cria relatório a partir de arquivo ou diretório"""
        if not os.path.exists(path):
            print(f"Caminho não encontrado: {path}")
            return

        # 1. Carregar imagens
        _safe_print(f"Processando: {path}")
        pil_images = self._load_images_from_path(path)

        if not pil_images:
            _safe_print("Nenhuma imagem encontrada.")
            return

        _safe_print(f"{len(pil_images)} imagem(ns) carregada(s)")

        # 2. Análise da IA
        ai_text_result = self.ai_analyzer.generate_diagnosis(pil_images)

        # 3. Geração do documento
        self._build_docx(path, pil_images, ai_text_result)

    def _build_docx(self, original_path, images, ai_text):
        doc = Document()
        self._setup_styles(doc)

        # Cabeçalho
        self._create_header(doc)

        # Seção de Imagens
        doc.add_heading('Imagens do Exame', level=1)
        for i, img in enumerate(images):
            # Converter PIL para stream bytes para o python-docx
            img_byte_arr = io.BytesIO()
            img.save(img_byte_arr, format='PNG')
            img_byte_arr.seek(0)

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(img_byte_arr, width=Inches(IMAGE_WIDTH_INCHES))
            doc.add_paragraph(f"Imagem {i + 1}").style = "Caption"

        # Seção do Laudo da IA
        doc.add_page_break()
        doc.add_heading('Laudo Sugerido (Gerado por IA)', level=1)

        # Nota de aviso
        warning_text = (
            "Nota: Este texto foi gerado automaticamente. "
            "O Médico Veterinário deve revisar e validar todas as "
            "informações."
        )
        warning = doc.add_paragraph(warning_text)
        warning.runs[0].font.color.rgb = RGBColor(255, 0, 0)
        warning.runs[0].font.italic = True

        # Inserir o texto da IA
        # O texto vem em Markdown, vamos apenas inseri-lo como texto
        # limpo por enquanto
        doc.add_paragraph(ai_text)

        # Rodapé de Assinatura
        doc.add_paragraph("_" * 30)
        doc.add_paragraph(
            "Assinatura e Carimbo do Veterinário Responsável"
        )

        # Salvar
        base_name = os.path.splitext(
            os.path.basename(original_path)
        )[0]
        filename = os.path.join(
            self.output_dir, f"Laudo_AI_{base_name}.docx"
        )
        doc.save(filename)
        _safe_print(f"Laudo gerado com sucesso: {filename}")

    def _setup_styles(self, doc):
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

    def _create_header(self, doc):
        head = doc.add_heading('LAUDO VETERINÁRIO DE IMAGEM', 0)
        head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table = doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        table.rows[0].cells[0].text = "Paciente/Tutor: ____________________"
        table.rows[0].cells[1].text = "Data: ____/____/______"
        doc.add_paragraph()


# --- Execução ---
if __name__ == "__main__":
    if len(sys.argv) < 2:
        _safe_print("Uso: python main.py <caminho_para_arquivo_ou_diretorio>")
        _safe_print("\nExemplos:")
        _safe_print("  python main.py exame.pdf")
        _safe_print("  python main.py imagem.jpg")
        _safe_print("  python main.py exame.dcm")
        _safe_print("  python main.py pasta_com_imagens/")
        sys.exit(1)

    path = sys.argv[1]
    generator = VetReportGenerator()
    api_key = os.getenv("GOOGLE_API_KEY", "SUA_API_KEY_AQUI")
    if not api_key or api_key == "SUA_API_KEY_AQUI":
        _safe_print("AVISO: Configure a GOOGLE_API_KEY no .env ou variável de ambiente.")
    else:
        generator.create_report(path)
