"""
PAICS - Aplicação Streamlit para Análise de Imagens Veterinárias com IA
Interface web para upload de PDFs, geração de laudos e edição.
"""

import streamlit as st
import fitz  # PyMuPDF
import google.generativeai as genai
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import io
import os
import tempfile
from dotenv import load_dotenv
from main import VetAIAnalyzer, VetReportGenerator
from fpdf import FPDF
from fpdf.enums import Align
import pytesseract
import re

# --- Carregar variáveis de ambiente do arquivo .env ---
load_dotenv()

# Configuração da página
st.set_page_config(
    page_title="PAICS - Análise de Imagens Veterinárias",
    page_icon="🐾",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Verificar e configurar API Key
# A chave pode ser definida no arquivo .env ou como variável de
# ambiente do sistema
API_KEY = os.getenv("GOOGLE_API_KEY", "SUA_API_KEY_AQUI")
if "SUA_API_KEY_AQUI" not in API_KEY and API_KEY:
    genai.configure(api_key=API_KEY)

# Inicializar o gerador de relatórios


@st.cache_resource
def init_report_generator():
    """Inicializa o gerador de relatórios (cached para melhor performance)"""
    # Usar diretório temporário do Streamlit ou variável de ambiente
    temp_dir = os.getenv("STREAMLIT_TEMP_DIR", "temp_laudos")
    return VetReportGenerator(output_dir=temp_dir)

# Inicializar o analisador de IA


@st.cache_resource
def init_ai_analyzer():
    """Inicializa o analisador de IA (cached para melhor performance)"""
    if "SUA_API_KEY_AQUI" in API_KEY or not API_KEY:
        return None
    return VetAIAnalyzer()


def pdf_to_images(pdf_bytes):
    """Converte PDF em bytes para lista de imagens PIL"""
    # Criar arquivo temporário
    with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
        tmp_file.write(pdf_bytes)
        tmp_path = tmp_file.name

    try:
        # Abrir PDF e converter páginas em imagens
        doc = fitz.open(tmp_path)
        pil_images = []

        # Usar zoom configurável via variável de ambiente
        zoom = float(os.getenv("PDF_ZOOM_FACTOR", "2.0"))

        for page in doc:
            # Zoom configurável para garantir que a IA veja detalhes finos
            pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom))
            mode = "RGBA" if pix.alpha else "RGB"
            img = Image.frombytes(mode, [pix.width, pix.height], pix.samples)
            pil_images.append(img)

        doc.close()
        return pil_images
    finally:
        # Limpar arquivo temporário
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


def create_docx_from_edited(images, edited_text, metadata=None):
    """Cria documento Word com laudo primeiro e depois as imagens"""
    doc = Document()

    # Estilos (configuráveis)
    style = doc.styles['Normal']
    style.font.name = os.getenv("DOC_FONT_NAME", "Calibri")
    font_size = int(os.getenv("DOC_FONT_SIZE", "11"))
    style.font.size = Pt(font_size)

    # Cabeçalho
    head = doc.add_heading('LAUDO VETERINÁRIO DE IMAGEM', 0)
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Tabela com informações do exame (4 linhas)
    if metadata:
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'

        # Linha 1: Paciente
        table.rows[0].cells[0].text = f"Paciente: {metadata.get('paciente', '____________________')}"
        table.rows[0].cells[1].text = f"Data: {metadata.get('data', '____/____/______')}"

        # Linha 2: Tutor
        table.rows[1].cells[0].text = f"Tutor: {metadata.get('tutor', '____________________')}"
        table.rows[1].cells[1].text = ""

        # Linha 3: Clínica
        table.rows[2].cells[0].text = f"Clínica: {metadata.get('clinica', '____________________')}"
        table.rows[2].cells[1].text = ""
    else:
        # Fallback para formato antigo
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.rows[0].cells[0].text = "Paciente/Tutor: ____________________"
        table.rows[0].cells[1].text = "Data: ____/____/______"

    doc.add_paragraph()

    # NOVA ORDEM: Laudo primeiro
    doc.add_heading('Laudo', level=1)

    # Nota de aviso (se foi gerado por IA)
    if edited_text and st.session_state.get('original_laudo'):
        warning_text = (
            "Nota: Este laudo foi gerado automaticamente e revisado pelo "
            "Médico Veterinário."
        )
        warning = doc.add_paragraph(warning_text)
        warning.runs[0].font.color.rgb = RGBColor(255, 140, 0)
        warning.runs[0].font.italic = True
        doc.add_paragraph()

    # Inserir o texto editado
    paragraphs = edited_text.split('\n')
    for para in paragraphs:
        if para.strip():
            doc.add_paragraph(para.strip())

    # Imagens depois do laudo
    doc.add_page_break()
    for i, img in enumerate(images):
        img_byte_arr = io.BytesIO()
        img.save(img_byte_arr, format='PNG')
        img_byte_arr.seek(0)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_width = float(os.getenv("IMAGE_WIDTH_INCHES", "4.5"))
        p.add_run().add_picture(img_byte_arr, width=Inches(image_width))
        doc.add_paragraph(f"Imagem {i + 1}").style = "Caption"

    # Rodapé de Assinatura
    doc.add_paragraph()
    doc.add_paragraph("_" * 50)
    doc.add_paragraph()
    doc.add_paragraph("Dra. Laís Costa Muchiutti")
    doc.add_paragraph("Médica Veterinária - CRMV-XX XXXXX")

    # Converter para bytes
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    return docx_bytes.getvalue()


class PDF(FPDF):
    """Classe customizada para PDF. O rodapé é adicionado manualmente."""

    def header(self):
        pass

    def footer(self):
        pass


def create_pdf_from_edited(images, edited_text, metadata=None):
    """Cria documento PDF com laudo primeiro, depois as imagens."""
    pdf = PDF('P', 'mm', 'A4')
    pdf.set_auto_page_break(auto=True, margin=15)

    def clean_unicode_text(text):
        """Remove ou substitui caracteres Unicode problemáticos para FPDF."""
        replacements = {
            ''': "'", ''': "'", '"': '"', '"': '"',
            '—': '-', '–': '-', '…': '...', '°': ' graus',
            '™': '', '®': '', '©': '', '•': '-',
            '→': '->', '←': '<-', '↑': '^', '↓': 'v',
            '≥': '>=', '≤': '<=', '≠': '!=', '×': 'x', '÷': '/',
        }
        for old, new in replacements.items():
            text = text.replace(old, new)
        text = text.replace('**', '')
        try:
            text.encode('latin-1')
        except UnicodeEncodeError:
            import unicodedata
            text = unicodedata.normalize('NFKD', text)
            text = text.encode('latin-1', 'ignore').decode('latin-1')
        return text

    pdf.add_page()
    pdf.set_font('Arial', 'B', 14)
    pdf.cell(0, 10, 'LAUDO VETERINARIO DE IMAGEM', ln=1, align=Align.C)
    pdf.ln(5)

    # Adicionar informações do exame
    if metadata:
        pdf.set_font('Arial', '', 10)
        pdf.cell(
            0, 6, f"Paciente: {clean_unicode_text(metadata.get('paciente', '____________________'))}", ln=1)
        pdf.cell(
            0, 6, f"Tutor: {clean_unicode_text(metadata.get('tutor', '____________________'))}", ln=1)
        pdf.cell(
            0, 6, f"Clinica: {clean_unicode_text(metadata.get('clinica', '____________________'))}", ln=1)
        pdf.cell(
            0, 6, f"Data: {clean_unicode_text(metadata.get('data', '____/____/______'))}", ln=1)
        pdf.ln(4)

    pdf.set_font('Arial', 'B', 12)
    pdf.cell(0, 10, 'Laudo', ln=1)
    pdf.ln(2)

    pdf.set_font('Arial', '', 11)
    clean_text = clean_unicode_text(edited_text)
    pdf.multi_cell(0, 5, clean_text)

    # Imagens
    pdf.add_page()
    max_width_mm = 180

    for i, img in enumerate(images):
        w_px, h_px = img.size
        aspect_ratio = h_px / w_px
        img_width_mm = max_width_mm
        img_height_mm = img_width_mm * aspect_ratio

        if pdf.get_y() + img_height_mm > (297 - 30):
            pdf.add_page()

        try:
            pdf.image(img, w=img_width_mm, h=img_height_mm)
        except Exception:
            buf = io.BytesIO()
            img.save(buf, format='PNG')
            buf.seek(0)
            pdf.image(buf, w=img_width_mm, h=img_height_mm)
            buf.close()

        pdf.set_font('Arial', 'I', 9)
        pdf.cell(0, 6, f"Imagem {i + 1}", ln=1, align=Align.C)
        pdf.ln(4)

    # Rodapé de assinatura
    pdf.set_y(-35)
    pdf.set_font('Arial', '', 10)
    pdf.cell(0, 10, "_" * 60, ln=1, align=Align.L)
    pdf.ln(2)
    pdf.set_font('Arial', 'B', 10)
    pdf.cell(0, 5, "Dra. Lais Costa Muchiutti", ln=1, align=Align.L)
    pdf.ln(2)
    pdf.set_font('Arial', '', 9)
    pdf.cell(0, 5, "Medica Veterinaria - CRMV-XX XXXXX", ln=1, align=Align.L)

    output = pdf.output(dest='S')
    if isinstance(output, bytearray):
        return bytes(output)
    if isinstance(output, bytes):
        return output
    if isinstance(output, str):
        return output.encode('latin-1', errors='replace')
    return bytes(output)


# configurar caminho do tesseract pelo env (opcional)
TESSERACT_CMD = os.getenv("TESSERACT_CMD", "")
if TESSERACT_CMD:
    pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD

# idioma para o tesseract (por + eng por padrão; ajuste se desejar)
TESSERACT_LANG = os.getenv("TESSERACT_LANG", "por+eng")


def ocr_extract_metadata(images: list) -> dict:
    """
    Roda OCR nas imagens e tenta extrair informações do cabeçalho:
      - Paciente, Tutor, Clínica Veterinária, Data do Exame
    Retorna dict com as informações encontradas.
    """
    # Padrões de busca
    name_patterns = [
        r'(?i)Paciente[:\s]*([A-ZÁÉÍÓÚÂÊÔÃÕÇ0-9\-\.\s,]{3,60})',
        r'(?i)Nome[:\s]*([A-ZÁÉÍÓÚÂÊÔÃÕÇ0-9\-\.\s,]{3,60})',
    ]
    tutor_patterns = [
        r'(?i)Tutor[:\s]*([A-ZÁÉÍÓÚÂÊÔÃÕÇ0-9\-\.\s,]{3,60})',
        r'(?i)Propriet[áa]rio[:\s]*([A-ZÁÉÍÓÚÂÊÔÃÕÇ0-9\-\.\s,]{3,60})',
        r'(?i)Respons[áa]vel[:\s]*([A-ZÁÉÍÓÚÂÊÔÃÕÇ0-9\-\.\s,]{3,60})',
    ]
    clinic_patterns = [
        r'(?i)Cl[ií]nica[:\s]*([A-ZÁÉÍÓÚÂÊÔÃÕÇ0-9\-\.\s,]{3,80})',
        r'(?i)Hospital[:\s]*([A-ZÁÉÍÓÚÂÊÔÃÕÇ0-9\-\.\s,]{3,80})',
        r'(?i)Centro[:\s]*([A-ZÁÉÍÓÚÂÊÔÃÕÇ0-9\-\.\s,]{3,80})',
    ]

    date_regex = re.compile(r'(\d{2}[\/\-]\d{2}[\/\-]\d{2,4}|\d{4}[\/\-]\d{2}[\/\-]\d{2})')
    month_names = r'(janeiro|fevereiro|março|marco|abril|maio|junho|julho|agosto|setembro|outubro|novembro|dezembro)'
    month_regex = re.compile(r'(\d{1,2}\s+' + month_names + r'\s+\d{4})', re.IGNORECASE)

    metadata = {
        'paciente': "",
        'tutor': "",
        'clinica': "",
        'data': ""
    }

    # Processar apenas as primeiras 2 imagens (cabeçalho geralmente está na primeira)
    for img in images[:2]:
        try:
            gray = img.convert('L')
            max_dim = 1600
            if max(gray.size) < max_dim:
                scale = max_dim / max(gray.size)
                new_size = (int(gray.width * scale), int(gray.height * scale))
                gray = gray.resize(new_size)

            text = pytesseract.image_to_string(gray, lang=TESSERACT_LANG)
            if not text or not text.strip():
                continue

            # Procurar por data
            if not metadata['data']:
                date_match = date_regex.search(text)
                if date_match:
                    metadata['data'] = date_match.group(1).strip()
                else:
                    month_match = month_regex.search(text)
                    if month_match:
                        metadata['data'] = month_match.group(1).strip()

            # Procurar por paciente
            if not metadata['paciente']:
                for pat in name_patterns:
                    m = re.search(pat, text)
                    if m:
                        candidate = m.group(1).strip()
                        candidate = re.sub(r'[^A-Za-zÀ-ÖØ-öø-ÿ0-9\s\-,\.]', '', candidate)
                        if 3 <= len(candidate) <= 60:
                            metadata['paciente'] = candidate
                            break

            # Procurar por tutor
            if not metadata['tutor']:
                for pat in tutor_patterns:
                    m = re.search(pat, text)
                    if m:
                        candidate = m.group(1).strip()
                        candidate = re.sub(r'[^A-Za-zÀ-ÖØ-öø-ÿ0-9\s\-,\.]', '', candidate)
                        if 3 <= len(candidate) <= 60:
                            metadata['tutor'] = candidate
                            break

            # Procurar por clínica
            if not metadata['clinica']:
                for pat in clinic_patterns:
                    m = re.search(pat, text)
                    if m:
                        candidate = m.group(1).strip()
                        candidate = re.sub(r'[^A-Za-zÀ-ÖØ-öø-ÿ0-9\s\-,\.]', '', candidate)
                        if 3 <= len(candidate) <= 80:
                            metadata['clinica'] = candidate
                            break

            # Se já encontrou todos, quebra
            if all(metadata.values()):
                break
        except Exception:
            continue

    # Normalizações
    if metadata['data']:
        metadata['data'] = metadata['data'].replace('\\', '/').strip()
    for key in metadata:
        if metadata[key]:
            metadata[key] = re.sub(r'\s{2,}', ' ', metadata[key]).strip()

    return metadata


def is_template_page(img: Image.Image) -> bool:
    """Retorna True se a imagem parecer um formulário/template vazio."""
    try:
        gray = img.convert('L')
        max_dim = 800
        if max(gray.size) > max_dim:
            scale = max_dim / max(gray.size)
            new_size = (int(gray.width * scale), int(gray.height * scale))
            gray = gray.resize(new_size)

        text = pytesseract.image_to_string(gray, lang=TESSERACT_LANG)
        if not text:
            return False

        t = text.lower()

        # Verificar se tem cabeçalho de formulário E campos vazios (underscores)
        has_template_header = "laudo veterinário de imagem" in t or "laudo veterinario de imagem" in t
        has_form_fields = "paciente/tutor" in t or "paciente / tutor" in t
        has_empty_underscores = text.count("___") >= 3  # Pelo menos 3 campos vazios

        # É template apenas se tiver TODOS os critérios
        return has_template_header and has_form_fields and has_empty_underscores
    except Exception:
        return False


def load_dicom_image(dicom_file) -> Image.Image:
    """Converte arquivo DICOM em imagem PIL"""
    try:
        import pydicom
        from pydicom.pixel_data_handlers.util import apply_voi_lut

        # Ler arquivo DICOM
        dicom = pydicom.dcmread(dicom_file)

        # Extrair dados de pixel
        data = apply_voi_lut(dicom.pixel_array, dicom)

        # Normalizar para 0-255
        data = data - data.min()
        data = data / data.max()
        data = (data * 255).astype('uint8')

        # Converter para PIL
        img = Image.fromarray(data)

        # Converter para RGB se necessário
        if img.mode != 'RGB':
            img = img.convert('RGB')

        return img
    except Exception as e:
        st.error(f"Erro ao processar DICOM: {e}")
        return None


def load_image_file(file) -> Image.Image:
    """Carrega arquivo de imagem (JPG, PNG, etc)"""
    try:
        img = Image.open(file)

        # Converter para RGB se necessário
        if img.mode != 'RGB':
            img = img.convert('RGB')

        return img
    except Exception as e:
        st.error(f"Erro ao processar imagem: {e}")
        return None


def process_uploaded_files(uploaded_files):
    """Processa múltiplos arquivos de imagem ou PDF"""
    all_images = []

    for file in uploaded_files:
        file_extension = file.name.lower().split('.')[-1]

        if file_extension == 'pdf':
            # Processar PDF
            pdf_bytes = file.read()
            pdf_images = pdf_to_images(pdf_bytes)
            all_images.extend(pdf_images)

        elif file_extension in ['dcm', 'dicom']:
            # Processar DICOM
            img = load_dicom_image(file)
            if img:
                all_images.append(img)

        elif file_extension in ['jpg', 'jpeg', 'png', 'bmp', 'tiff']:
            # Processar imagem comum
            img = load_image_file(file)
            if img:
                all_images.append(img)
        else:
            st.warning(f"⚠️ Formato não suportado: {file.name}")

    return all_images


# Sidebar com configurações
with st.sidebar:
    st.title("⚙️ Configurações")

    # Verificação de API Key
    if "SUA_API_KEY_AQUI" in API_KEY or not API_KEY:
        st.error("⚠️ API Key não configurada!")
        st.info("Configure a variável de ambiente GOOGLE_API_KEY")
        st.code("export GOOGLE_API_KEY='sua_chave_aqui'")
        api_configured = False
    else:
        st.success("✅ API Key configurada")
        api_configured = True

    st.divider()

    st.subheader("ℹ️ Sobre")
    st.info(
        """
        **PAICS** - Sistema de Análise de Imagens Veterinárias com IA

        Faça upload de um PDF com imagens de raio-x ou ultrassom
        veterinário e receba um laudo técnico gerado por IA.

        O laudo gerado deve ser sempre revisado e validado por um Médico
        Veterinário qualificado.
        """
    )

# Título principal
st.title("🐾 PAICS - Análise de Imagens Veterinárias com IA")
st.markdown("Sistema automatizado para geração de laudos técnicos de imagens veterinárias")

# Inicializar variáveis de sessão
if 'pdf_uploaded' not in st.session_state:
    st.session_state.pdf_uploaded = False
if 'images' not in st.session_state:
    st.session_state.images = []
if 'images_for_report' not in st.session_state:  # NOVO
    st.session_state.images_for_report = []
if 'laudo_gerado' not in st.session_state:
    st.session_state.laudo_gerado = False
if 'laudo_text' not in st.session_state:
    st.session_state.laudo_text = ""
if 'original_laudo' not in st.session_state:
    st.session_state.original_laudo = ""
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []
if 'ocr_paciente' not in st.session_state:
    st.session_state['ocr_paciente'] = ""
if 'ocr_data' not in st.session_state:
    st.session_state['ocr_data'] = ""
if 'ocr_raw' not in st.session_state:
    st.session_state['ocr_raw'] = ""
if 'ocr_done' not in st.session_state:
    st.session_state['ocr_done'] = False
if 'paciente_nome' not in st.session_state:
    st.session_state['paciente_nome'] = ""
if 'data_exame' not in st.session_state:
    st.session_state['data_exame'] = ""
if 'metadata' not in st.session_state:
    st.session_state.metadata = {}

# NOVA SEÇÃO - Upload de Imagens apenas
st.header("📤 Upload de Imagens do Exame")

uploaded_files = st.file_uploader(
    "Selecione as imagens do exame veterinário",
    type=['jpg', 'jpeg', 'png', 'dcm', 'dicom', 'bmp', 'tiff'],
    accept_multiple_files=True,
    help="Faça upload de imagens JPG/PNG ou arquivos DICOM de raio-x/ultrassom"
)

if uploaded_files:
    # Processar arquivos
    with st.spinner("Processando arquivos..."):
        all_images = process_uploaded_files(uploaded_files)
        total_images = len(all_images)

        if total_images == 0:
            st.warning("⚠️ Nenhuma imagem válida encontrada nos arquivos.")
            st.session_state.images = []
            st.session_state.images_for_report = []
            st.session_state.metadata = {}
        else:
            # Todas as imagens para análise
            st.session_state.images = all_images
            st.session_state.images_for_report = all_images

            # Extrair metadata das imagens
            with st.spinner("Extraindo informações do cabeçalho..."):
                st.session_state.metadata = ocr_extract_metadata(all_images)

            st.success(f"✅ {total_images} imagem(ns) carregada(s) com sucesso!")

            # Mostrar informações extraídas
            if any(st.session_state.metadata.values()):
                with st.expander("ℹ️ Informações Extraídas", expanded=False):
                    col1, col2 = st.columns(2)
                    with col1:
                        if st.session_state.metadata.get('paciente'):
                            st.write(f"**Paciente:** {st.session_state.metadata['paciente']}")
                        if st.session_state.metadata.get('tutor'):
                            st.write(f"**Tutor:** {st.session_state.metadata['tutor']}")
                    with col2:
                        if st.session_state.metadata.get('clinica'):
                            st.write(f"**Clínica:** {st.session_state.metadata['clinica']}")
                        if st.session_state.metadata.get('data'):
                            st.write(f"**Data:** {st.session_state.metadata['data']}")

# Mensagem se nenhum arquivo foi carregado
if not uploaded_files:
    st.info("👆 Faça upload das imagens do exame para começar")

# Seção de geração de laudo
if st.session_state.get('images') and api_configured:
    col1, col2, col3 = st.columns([1, 1, 1])
    with col2:
        generate_button = st.button("🚀 Gerar Laudo", type="primary",
                                    use_container_width=True)

    if generate_button or st.session_state.laudo_gerado:
        if generate_button:
            # Limpar chat anterior ao gerar novo laudo
            st.session_state.chat_history = []
            # Gerar laudo
            with st.spinner("🤖 Analisando imagens com IA..."):
                try:
                    ai_analyzer = init_ai_analyzer()
                    if ai_analyzer:
                        laudo = ai_analyzer.generate_diagnosis(st.session_state.images)
                        st.session_state.laudo_text = laudo
                        st.session_state.original_laudo = laudo
                        st.session_state.laudo_gerado = True
                        st.rerun()
                    else:
                        st.error("Erro ao inicializar o analisador de IA")
                except Exception as e:
                    st.error(f"Erro ao gerar laudo: {str(e)}")
                    st.exception(e)

if st.session_state.laudo_gerado and st.session_state.laudo_text:
    st.header("📝 Laudo Gerado - Edite conforme necessário")

    # Editor de texto
    edited_laudo = st.text_area(
        "Laudo (edite o texto gerado pela IA)",
        value=st.session_state.laudo_text,
        height=400
    )

    # Atualizar laudo editado na sessão
    st.session_state.laudo_text = edited_laudo

    # Botões de ação
    col1, col2, col3, col4 = st.columns(4)

    with col1:
        if st.button("💾 Salvar Edições", use_container_width=True):
            st.session_state.laudo_text = edited_laudo
            st.success("✅ Edições salvas!")

    with col2:
        # Botão para gerar novo laudo
        if st.button("🔄 Gerar Novo Laudo", use_container_width=True):
            st.session_state.laudo_gerado = False
            st.session_state.laudo_text = ""
            st.session_state.chat_history = []
            st.rerun()

    # Preparar nome do arquivo para download
    paciente_nome_str = st.session_state.get('paciente_nome', "")

    # Gerar nome base do arquivo
    if uploaded_files and len(uploaded_files) > 0:
        # Usar nome do primeiro arquivo carregado
        first_file = uploaded_files[0]
        nome_base = os.path.splitext(first_file.name)[0]
    else:
        # Fallback se não houver arquivo
        nome_base = "Laudo"

    # Adicionar nome do paciente se disponível
    if paciente_nome_str:
        nome_arquivo = f"{paciente_nome_str}_{nome_base}"
    else:
        nome_arquivo = f"Laudo_{nome_base}"

     # Gerar documento Word para download
    docx_bytes = create_docx_from_edited(
        st.session_state.images_for_report,
        edited_laudo,
        st.session_state.get('metadata', {})
    )

    with col3:
        st.download_button(
            label="📥 Baixar como Word",
            data=docx_bytes,
            file_name=f"{nome_arquivo}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

    with col4:
        try:
            pdf_bytes = create_pdf_from_edited(
                st.session_state.images_for_report,
                edited_laudo,
                st.session_state.get('metadata', {})
            )
            st.download_button(
                label="📥 Baixar como PDF",
                data=pdf_bytes,
                file_name=f"{nome_arquivo}.pdf",
                mime="application/pdf",
                use_container_width=True,
                type="primary"
            )
        except Exception as e:
            st.error(f"Erro ao gerar PDF: {e}")
            st.exception(e)

    st.divider()

    # Chat com a IA sobre o laudo - VERSÃO MELHORADA COM FOCO NO USUÁRIO
    st.subheader("💬 Refinar Laudo com IA")
    st.markdown(
        "**Digite suas solicitações abaixo para editar o laudo. As alterações serão aplicadas automaticamente.**")

    # Input do usuário PRIMEIRO (prioridade)
    prompt = st.text_input(
        "Digite sua solicitação:",
        placeholder="Ex: 'Remova a menção sobre derrame pleural' ou 'Adicione mais detalhes sobre a coluna'",
        key="user_prompt_input"
    )

    cols_space, col_send, cols_space2 = st.columns([1, 1, 1])

    with col_send:
        send_button = st.button("🚀 Enviar", type="primary",
                                use_container_width=True)

    # Processar solicitação do usuário
    if send_button and prompt:
        with st.spinner("🤖 Processando sua solicitação e atualizando o laudo..."):
            ai_analyzer = init_ai_analyzer()
            if ai_analyzer:
                # Prompt otimizado para merge automático
                chat_prompt = f"""
                Você é um assistente especializado em radiologia veterinária que trabalha em colaboração com veterinários.
                
                LAUDO ATUAL:
                ---
                {edited_laudo}
                ---
                
                SOLICITAÇÃO DO VETERINÁRIO: "{prompt}"
                
                INSTRUÇÕES IMPORTANTES:
                1. Analise a solicitação cuidadosamente
                2. Se for para REMOVER: remova completamente e ajuste o texto para manter coesão
                3. Se for para ADICIONAR: integre naturalmente ao laudo, consultando as imagens se necessário
                4. Se for para ALTERAR: faça a alteração mantendo o restante intacto
                5. Se for REVISAR: corrija erros mantendo o conteúdo técnico
                6. Se for REORGANIZAR: reestruture mantendo todo o conteúdo
                7. SEMPRE retorne o LAUDO COMPLETO já editado (não apenas trechos)
                8. NÃO adicione explicações, comentários ou introduções
                9. Mantenha tom profissional e técnico veterinário
                10. Preserve formatação e estrutura
                
                RETORNE APENAS O LAUDO COMPLETO EDITADO, SEM NENHUM TEXTO ADICIONAL:
                """

                content_for_ai = [chat_prompt] + st.session_state.images

                try:
                    response = ai_analyzer.model.generate_content(content_for_ai)
                    response_text = response.text.strip()

                    # Limpar formatação markdown extra
                    if response_text.startswith("```"):
                        lines = response_text.split("\n")
                        response_text = "\n".join(lines[1:-1]) if len(lines) > 2 else response_text

                    # APLICAR AUTOMATICAMENTE as mudanças
                    st.session_state.laudo_text = response_text

                    # Registrar no histórico
                    st.session_state.chat_history.append({
                        "role": "user",
                        "content": prompt
                    })
                    st.session_state.chat_history.append({
                        "role": "assistant",
                        "content": "Laudo atualizado com sucesso!",
                        "updated_laudo": response_text
                    })

                    st.success(f"✅ Laudo atualizado! Solicitação: '{prompt[:50]}...'")
                    st.rerun()

                except Exception as e:
                    st.error(f"Erro ao processar solicitação: {str(e)}")
            else:
                st.error("Analisador de IA não inicializado.")

    st.divider()

    # Histórico de alterações (colapsado por padrão)
    with st.expander("📜 Histórico de Alterações", expanded=False):
        if st.session_state.chat_history:
            for idx, message in enumerate(st.session_state.chat_history):
                if message["role"] == "user":
                    st.markdown(f"**🔵 Você:** {message['content']}")
                else:
                    st.markdown(f"**🤖 IA:** {message['content']}")
                st.divider()
        else:
            st.info("Nenhuma alteração ainda. Faça uma solicitação acima para começar.")

    # Processar comando rápido pendente (se houver)
    if st.session_state.chat_history and st.session_state.chat_history[-1]["role"] == "user":
        last_msg = st.session_state.chat_history[-1]
        if last_msg.get("is_quick"):
            with st.spinner("🤖 Processando atalho rápido..."):
                ai_analyzer = init_ai_analyzer()
                if ai_analyzer:
                    chat_prompt = f"""
                    Você é um assistente especializado em radiologia veterinária.
                    
                    LAUDO ATUAL:
                    ---
                    {edited_laudo}
                    ---
                    
                    INSTRUÇÃO: {last_msg['content']}
                    
                    Retorne APENAS o laudo completo atualizado, sem explicações adicionais.
                    Mantenha tom profissional e técnico veterinário.
                    """
                    content_for_ai = [chat_prompt] + st.session_state.images

                    try:
                        response = ai_analyzer.model.generate_content(content_for_ai)
                        response_text = response.text.strip()

                        if response_text.startswith("```"):
                            response_text = "\n".join(response_text.split("\n")[1:-1])

                        # APLICAR AUTOMATICAMENTE
                        st.session_state.laudo_text = response_text

                        st.session_state.chat_history.append({
                            "role": "assistant",
                            "content": "Laudo atualizado!",
                            "updated_laudo": response_text
                        })

                        st.success("✅ Laudo atualizado com sucesso!")
                        st.rerun()
                    except Exception as e:
                        st.error(f"Erro: {str(e)}")
                else:
                    st.error("Analisador de IA não inicializado.")
